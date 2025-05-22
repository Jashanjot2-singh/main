import streamlit as st
st.set_page_config(page_title="📑 RAG Chatbot with OCR", layout="wide")  # FIRST command

from datetime import datetime
from PIL import Image
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import ChatPromptTemplate
from qdrant_client import QdrantClient
from qdrant_client.http.models import PointStruct, VectorParams, Distance
from docx import Document as DocxDocument
import tempfile
import os
import pandas as pd
import requests

# Load TrOCR
@st.cache_resource
def load_trocr():
    processor = TrOCRProcessor.from_pretrained("microsoft/trocr-base-handwritten")
    model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-base-handwritten")
    return processor, model

processor, trocr_model = load_trocr()

def ocr_image(image_path):
    image = Image.open(image_path).convert("RGB")
    pixel_values = processor(images=image, return_tensors="pt").pixel_values
    generated_ids = trocr_model.generate(pixel_values)
    return processor.batch_decode(generated_ids, skip_special_tokens=True)[0]

# Loaders
def load_pdf(path): return PyPDFLoader(path).load()
def load_txt(path): return [Document(page_content=open(path, encoding="utf-8").read())]
def load_docx(path):
    doc = DocxDocument(path)
    return [Document(page_content="\n".join(p.text for p in doc.paragraphs))]
def load_excel(path):
    df = pd.read_excel(path)
    return df

def to_docs_from_excel(df):
    return [Document(page_content=" ".join(map(str, row))) for _, row in df.iterrows()]

# Split Text
def split_text(docs, size=500, overlap=50):
    return RecursiveCharacterTextSplitter(chunk_size=size, chunk_overlap=overlap).split_documents(docs)

# Embeddings + Qdrant
def compute_embeddings(texts):
    embeddings = []
    for t in texts:
        r = requests.post("http://localhost:11434/api/embeddings", json={"model": "nomic-embed-text", "prompt": t})
        embeddings.append(r.json().get("embedding", []))
    return embeddings

def create_qdrant(name, dim):
    QdrantClient(host="localhost", port=6333).recreate_collection(name, vectors_config=VectorParams(size=dim, distance=Distance.COSINE))

def insert_qdrant(name, texts, vectors):
    client = QdrantClient(host="localhost", port=6333)
    pts = [PointStruct(id=i, vector=vectors[i], payload={"text": texts[i]}) for i in range(len(texts))]
    client.upsert(name, points=pts)

def search_qdrant(name, q_vector):
    client = QdrantClient(host="localhost", port=6333)
    results = client.search(collection_name=name, query_vector=q_vector, limit=10)
    return [r.payload.get("text", "") for r in results]

def run_rag(query, docs):
    chunks = split_text(docs)
    texts = [c.page_content for c in chunks]
    embeddings = compute_embeddings(texts)
    create_qdrant("query_collection", len(embeddings[0]))
    insert_qdrant("query_collection", texts, embeddings)
    query_embed = compute_embeddings([query])[0]
    context = " ".join(search_qdrant("query_collection", query_embed))

    prompt = ChatPromptTemplate.from_template("""
    Answer the question based only on the following context:
    {context}
    Question: {question}
    Provide a detailed and well-supported answer.
    """).format(context=context, question=query)

    response = requests.post("http://localhost:11434/api/generate", json={
        "model": "llama3.2:1b",
        "prompt": prompt,
        "stream": False
    })
    return response.json().get("response", "No response")

# State
if "files" not in st.session_state:
    st.session_state.files = []

if "query_results" not in st.session_state:
    st.session_state.query_results = {}

# Upload
st.title("📄 RAG Chatbot + OCR for Scanned Docs")
uploaded = st.file_uploader("Upload documents or images", type=["pdf", "txt", "docx", "xlsx", "png", "jpg", "jpeg"], accept_multiple_files=True)

if uploaded:
    existing_names = [f["name"] for f in st.session_state.files]
    for file in uploaded:
        if file.name in existing_names:
            continue  # skip duplicates

        ext = file.name.split(".")[-1].lower()
        path = os.path.join(tempfile.gettempdir(), file.name)
        with open(path, "wb") as f:
            f.write(file.read())

        # Extract content
        content = ""
        df_data = None
        if ext in ["png", "jpg", "jpeg"]:
            content = ocr_image(path)
        elif ext == "pdf":
            content = "\n".join([p.page_content for p in load_pdf(path)])
        elif ext == "txt":
            content = load_txt(path)[0].page_content
        elif ext == "docx":
            content = load_docx(path)[0].page_content
        elif ext == "xlsx":
            df_data = load_excel(path)
            content = "\n".join([" ".join(map(str, row)) for _, row in df_data.iterrows()])

        st.session_state.files.append({
            "name": file.name,
            "path": path,
            "text": content,
            "ext": ext,
            "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "df": df_data
        })

# File list + Actions
st.subheader("📂 Uploaded Files")
for i, file in enumerate(st.session_state.files):
    col1, col2 = st.columns([4, 6])
    col1.markdown(f"**📄 {file['name']}**  \n🕒 {file['date']}")

    with col2:
        view_button = st.button(f"👁 View {file['name']}", key=f"view_{i}")
        if view_button:
            with st.expander(f"Content of {file['name']}", expanded=True):
                if file["ext"] == "xlsx" and file["df"] is not None:
                    st.dataframe(file["df"], use_container_width=True)
                else:
                    st.code(file["text"][:4000] + ("..." if len(file["text"]) > 4000 else ""))

        with st.form(f"query_form_{i}"):
            query = st.text_input(f"❓ Enter your query for `{file['name']}`:", key=f"query_{i}")
            submitted = st.form_submit_button("Run Query")
            if submitted and query:
                if file["ext"] == "pdf":
                    docs = load_pdf(file["path"])
                elif file["ext"] == "txt":
                    docs = load_txt(file["path"])
                elif file["ext"] == "docx":
                    docs = load_docx(file["path"])
                elif file["ext"] == "xlsx":
                    docs = to_docs_from_excel(file["df"])
                elif file["ext"] in ["png", "jpg", "jpeg"]:
                    docs = [Document(page_content=file["text"])]
                else:
                    st.warning("Unsupported file type.")
                    continue

                result = run_rag(query, docs)
                st.session_state.query_results[file["name"]] = result

        # Show result (outside form)
        if file["name"] in st.session_state.query_results:
            st.success(f"💬 Answer: {st.session_state.query_results[file['name']]}")
